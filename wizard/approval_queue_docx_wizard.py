# -*- coding: utf-8 -*-
"""Generate a Word (.docx) or Excel (.xlsx) report of the Board or Floor
approval queue, for the Secretary to paste into meeting minutes or share
as a spreadsheet the Board can sort/filter.

Word layout — copy-paste-friendly for minutes:
  - Title line, date + counts, motion header
  - Numbered list of requisitions
  - Grand total, vote-recording block

Excel layout — one row per requisition with an editable Vote column
  - Filter/sort by dept, amount, etc.
  - Vote column ready for tally marks

Files named e.g. "Board Approval Queue - 2026-08-16.docx" / ".xlsx".
"""
import base64
import io
from datetime import date

from odoo import _, api, fields, models
from odoo.exceptions import UserError

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None  # module install guard; raise at runtime

# openpyxl is a hard dependency of Odoo (used by base import/export)
# so we don't need a manifest external_dependencies entry for it.
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter


QUEUE_TYPES = [
    ('board', 'Board Approval Queue'),
    ('floor', 'Floor Approval Queue'),
]


class ApprovalQueueDocxWizard(models.TransientModel):
    _name = "elks.approval.queue.docx.wizard"
    _description = "Approval Queue — Word / Excel Export"

    queue_type = fields.Selection(
        QUEUE_TYPES, string="Queue", required=True, default='board',
    )
    output_format = fields.Selection([
        ('docx', 'Word (.docx)'),
        ('xlsx', 'Excel (.xlsx)'),
    ], string="Format", required=True, default='docx')

    def action_generate(self):
        """Dispatcher — Word or Excel based on output_format."""
        self.ensure_one()
        if self.output_format == 'xlsx':
            return self.action_generate_xlsx()
        return self.action_generate_docx()

    def _get_queue_pos(self):
        pos = self.env['purchase.order'].search(
            [('x_approval_state', '=', self.queue_type)],
            order='x_requesting_department_id, name',
        )
        if not pos:
            raise UserError(_(
                "No requisitions are currently in the %s queue.",
                dict(QUEUE_TYPES)[self.queue_type],
            ))
        return pos

    def _save_and_download(self, data: bytes, ext: str, mimetype: str):
        today_str = date.today().strftime("%Y-%m-%d")
        queue_label = dict(QUEUE_TYPES)[self.queue_type]
        fname = f"{queue_label} - {today_str}.{ext}"
        att = self.env['ir.attachment'].create({
            'name': fname,
            'datas': base64.b64encode(data),
            'res_model': self._name,
            'res_id': self.id,
            'mimetype': mimetype,
        })
        return {
            'type': 'ir.actions.act_url',
            'url': f'/web/content/{att.id}?download=true',
            'target': 'self',
        }

    def action_generate_docx(self):
        """Build the .docx, save as ir.attachment, return download URL."""
        self.ensure_one()
        if Document is None:
            raise UserError(_(
                "python-docx is not installed on the server. "
                "Ask your admin to run `pip install python-docx`."
            ))
        pos = self._get_queue_pos()
        doc = self._build_docx(pos)
        buf = io.BytesIO()
        doc.save(buf)
        return self._save_and_download(
            buf.getvalue(),
            ext="docx",
            mimetype=(
                'application/vnd.openxmlformats-officedocument.'
                'wordprocessingml.document'
            ),
        )

    def action_generate_xlsx(self):
        """Build an .xlsx workbook of the queue and download it."""
        self.ensure_one()
        pos = self._get_queue_pos()
        wb = self._build_xlsx(pos)
        buf = io.BytesIO()
        wb.save(buf)
        return self._save_and_download(
            buf.getvalue(),
            ext="xlsx",
            mimetype=(
                'application/vnd.openxmlformats-officedocument.'
                'spreadsheetml.sheet'
            ),
        )

    # ------------------------------------------------------------------
    # DOCX layout
    # ------------------------------------------------------------------
    def _build_docx(self, pos):
        queue_label = dict(QUEUE_TYPES)[self.queue_type]
        stage = "Board" if self.queue_type == 'board' else "Floor"

        doc = Document()

        # Global font
        style = doc.styles['Normal']
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        # ── Title ────────────────────────────────────────────────
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(queue_label)
        run.bold = True
        run.font.size = Pt(16)

        # ── Date + summary line ─────────────────────────────────
        subtotal = sum(p.amount_total for p in pos)
        over_budget = len(pos.filtered('x_has_over_budget_lines'))
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub.add_run(
            f"Prepared {date.today().strftime('%B %d, %Y')}  "
            f"·  {len(pos)} requisition(s)  ·  Total: ${subtotal:,.2f}"
        )
        sub_run.italic = True
        sub_run.font.size = Pt(10)
        if over_budget:
            warn = doc.add_paragraph()
            warn.alignment = WD_ALIGN_PARAGRAPH.CENTER
            wr = warn.add_run(f"⚠ {over_budget} requisition(s) over budget")
            wr.italic = True
            wr.font.size = Pt(10)
            wr.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

        doc.add_paragraph()  # spacer

        # ── Motion header ───────────────────────────────────────
        motion = doc.add_paragraph()
        m_run = motion.add_run(
            f"MOTION: To approve the following requisitions "
            f"as presented to the {stage}:"
        )
        m_run.bold = True

        # ── Requisition list ────────────────────────────────────
        for idx, po in enumerate(pos, start=1):
            p = doc.add_paragraph(style='List Number')
            # Line 1: PO # + title + vendor + amount (single paragraph)
            title_part = po.x_requisition_title or "Untitled"
            po_run = p.add_run(f"{po.name} — ")
            po_run.bold = True
            p.add_run(f"{title_part}  ·  ")
            vendor_run = p.add_run(f"{po.partner_id.display_name or '—'}")
            vendor_run.italic = True
            p.add_run("  ·  ")
            amt_run = p.add_run(f"${po.amount_total:,.2f}")
            amt_run.bold = True

            # Line 2: department / committee (indented small text)
            meta_bits = []
            if po.x_requesting_department_id:
                meta_bits.append(f"Dept: {po.x_requesting_department_id.display_name}")
            if po.x_requesting_committee_id:
                meta_bits.append(f"Committee: {po.x_requesting_committee_id.display_name}")
            if po.x_has_over_budget_lines:
                meta_bits.append("⚠ OVER BUDGET")
            if meta_bits:
                meta = doc.add_paragraph()
                meta.paragraph_format.left_indent = Inches(0.5)
                meta_run = meta.add_run("    " + " · ".join(meta_bits))
                meta_run.italic = True
                meta_run.font.size = Pt(9)
                if po.x_has_over_budget_lines:
                    meta_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

        # ── Grand total ─────────────────────────────────────────
        doc.add_paragraph()
        gt = doc.add_paragraph()
        gt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        gt_run = gt.add_run(f"GRAND TOTAL: ${subtotal:,.2f}")
        gt_run.bold = True
        gt_run.font.size = Pt(12)

        # ── Vote-recording block ────────────────────────────────
        doc.add_paragraph()
        vote_intro = doc.add_paragraph()
        vi_run = vote_intro.add_run("Motion by: ______________________   Second by: ______________________")
        vi_run.font.size = Pt(11)

        vote_row = doc.add_paragraph()
        vr_run = vote_row.add_run(
            "Vote:    YES ______    NO ______    ABSTAIN ______    "
            "Motion:  ☐ Carried    ☐ Failed"
        )
        vr_run.font.size = Pt(11)

        doc.add_paragraph()  # spacer

        # ── Signature line ─────────────────────────────────────
        sig = doc.add_paragraph()
        sig_run = sig.add_run(
            f"{stage} Chair Signature: _________________________________   "
            f"Date: __________"
        )
        sig_run.font.size = Pt(10)

        return doc

    # ------------------------------------------------------------------
    # XLSX layout — meeting worksheet the Board can sort/filter/tally in
    # ------------------------------------------------------------------
    def _build_xlsx(self, pos):
        queue_label = dict(QUEUE_TYPES)[self.queue_type]
        stage = "Board" if self.queue_type == 'board' else "Floor"
        subtotal = sum(p.amount_total for p in pos)
        over_budget = len(pos.filtered('x_has_over_budget_lines'))

        wb = Workbook()
        ws = wb.active
        ws.title = f"{stage} Queue"

        # ── Styles ──────────────────────────────────────────
        title_font  = Font(name="Calibri", bold=True, size=16, color="FFFFFF")
        title_fill  = PatternFill("solid", start_color="2C3E50")
        sub_font    = Font(name="Calibri", italic=True, size=10, color="555555")
        hdr_font    = Font(name="Calibri", bold=True, size=11, color="FFFFFF")
        hdr_fill    = PatternFill("solid", start_color="2C3E50")
        hdr_align   = Alignment(horizontal="center", vertical="center", wrap_text=True)
        body_font   = Font(name="Calibri", size=11)
        total_font  = Font(name="Calibri", bold=True, size=12)
        red_font    = Font(name="Calibri", bold=True, size=11, color="C00000")
        red_fill    = PatternFill("solid", start_color="FFF2F2")
        thin        = Side(style="thin", color="CCCCCC")
        box         = Border(left=thin, right=thin, top=thin, bottom=thin)

        # ── Row 1: Title (merged) ──────────────────────────
        ws.merge_cells("A1:F1")
        c = ws["A1"]
        c.value = queue_label
        c.font = title_font
        c.fill = title_fill
        c.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[1].height = 30

        # ── Row 2: Date + counts ───────────────────────────
        ws.merge_cells("A2:F2")
        c = ws["A2"]
        c.value = (
            f"Prepared {date.today().strftime('%B %d, %Y')}   ·   "
            f"{len(pos)} requisition(s)   ·   Total: ${subtotal:,.2f}"
            + (f"   ·   ⚠ {over_budget} over budget" if over_budget else "")
        )
        c.font = sub_font
        c.alignment = Alignment(horizontal="center")

        # ── Row 4: Column headers ──────────────────────────
        headers = ["PO #", "Requisition Title", "Vendor",
                   "Dept / Committee", "Amount", "Vote (Y/N/Abs)"]
        for col_idx, h in enumerate(headers, start=1):
            cell = ws.cell(row=4, column=col_idx, value=h)
            cell.font = hdr_font
            cell.fill = hdr_fill
            cell.alignment = hdr_align
            cell.border = box

        # ── Data rows ──────────────────────────────────────
        row = 5
        for po in pos:
            over = po.x_has_over_budget_lines
            row_font = red_font if over else body_font
            row_fill = red_fill if over else None

            values = [
                po.name,
                (po.x_requisition_title or "Untitled")
                    + (" ⚠ OVER BUDGET" if over else ""),
                po.partner_id.display_name or "—",
                (po.x_requesting_department_id.display_name or "")
                + (f"\n{po.x_requesting_committee_id.display_name}"
                   if po.x_requesting_committee_id else ""),
                po.amount_total,
                "",  # Vote — Secretary fills in
            ]
            for col_idx, val in enumerate(values, start=1):
                cell = ws.cell(row=row, column=col_idx, value=val)
                cell.font = row_font
                cell.border = box
                cell.alignment = Alignment(
                    wrap_text=True, vertical="top",
                    horizontal="right" if col_idx == 5 else "left",
                )
                if col_idx == 5:
                    cell.number_format = "$#,##0.00"
                if row_fill:
                    cell.fill = row_fill
            row += 1

        # ── Grand total row ────────────────────────────────
        total_row = row + 1
        ws.cell(row=total_row, column=4, value="GRAND TOTAL").font = total_font
        ws.cell(row=total_row, column=4).alignment = Alignment(horizontal="right")
        gt = ws.cell(row=total_row, column=5, value=f"=SUM(E5:E{row-1})")
        gt.font = total_font
        gt.number_format = "$#,##0.00"
        gt.border = box

        # ── Vote tally section ─────────────────────────────
        row = total_row + 3
        ws.cell(row=row, column=1, value="Motion by:").font = total_font
        ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=3)
        ws.cell(row=row, column=4, value="Second by:").font = total_font
        ws.merge_cells(start_row=row, start_column=5, end_row=row, end_column=6)
        row += 2
        ws.cell(row=row, column=1, value="Vote tally:").font = total_font
        ws.cell(row=row, column=2, value="YES:")
        ws.cell(row=row, column=3, value="NO:")
        ws.cell(row=row, column=4, value="ABSTAIN:")
        ws.cell(row=row, column=5, value="Motion:")
        ws.cell(row=row, column=6, value="☐ Carried  ☐ Failed")
        row += 2
        ws.cell(row=row, column=1,
                value=f"{stage} Chair Signature:").font = total_font
        ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=4)
        ws.cell(row=row, column=5, value="Date:").font = total_font

        # ── Column widths ──────────────────────────────────
        widths = [12, 42, 30, 28, 14, 22]
        for i, w in enumerate(widths, start=1):
            ws.column_dimensions[get_column_letter(i)].width = w

        # Freeze the header row so Board can scroll long queues
        ws.freeze_panes = "A5"

        # Auto-filter on the header row for easy sort by dept, amount, etc.
        ws.auto_filter.ref = f"A4:F{total_row - 2}"

        return wb
